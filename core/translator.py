import os
import re
from docx import Document
from PyPDF2 import PdfReader
import markdown
from langdetect import detect
import json
import logging
from typing import Dict, List, Optional, Tuple, Any

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class Translator:
    """翻译引擎"""
    
    def __init__(self, ai_engine, term_loader, settings):
        """初始化翻译引擎"""
        self.ai_engine = ai_engine
        self.term_loader = term_loader
        self.settings = settings
        
        # 初始化术语库
        self.terms = self._load_terms()
        
        # 确保模型可用于翻译
        self.model = ai_engine
        
        # 记录术语库加载状态
        print(f"术语库初始化状态：已加载 {len(self.terms)} 个术语")
        if len(self.terms) > 0:
            sample_terms = list(self.terms.items())[:3]
            for key, value in sample_terms:
                print(f"术语示例: {key} -> {value.get('target_term', '')}")
    
    def _load_terms(self) -> Dict:
        """从文件加载术语库"""
        try:
            terms_path = os.path.join('data', 'terms', 'terms.json')
            if os.path.exists(terms_path):
                with open(terms_path, 'r', encoding='utf-8') as f:
                    terms = json.load(f)
                    logger.info(f"成功加载术语库，共 {len(terms)} 个术语")
                    
                    # 确保AIEngine也加载了术语库
                    if hasattr(self, 'ai_engine') and self.ai_engine:
                        if not hasattr(self.ai_engine, 'terms') or not self.ai_engine.terms:
                            self.ai_engine.terms = terms
                            self.ai_engine.load_term_database()
                            logger.info("已将术语库同步到AI引擎")
                    
                    return terms
            else:
                logger.warning(f"术语库文件不存在: {terms_path}")
                return {}
        except Exception as e:
            logger.error(f"加载术语库出错: {e}")
            return {}
    
    def translate(self, text: str, use_termbase: bool = True, source_lang: Optional[str] = None, target_lang: str = "en", force_terms: bool = False, matched_terms: List = None) -> str:
        """
        翻译文本
        
        Args:
            text: 要翻译的文本
            use_termbase: 是否使用术语库
            source_lang: 源语言代码，如果为None则自动检测
            target_lang: 目标语言代码
            force_terms: 是否强制使用匹配到的术语
            matched_terms: 匹配到的术语列表，如果提供则直接使用
            
        Returns:
            翻译后的文本
        """
        if not text or not text.strip():
            logger.warning("翻译内容为空")
            return ""
            
        logger.info(f"开始翻译文本 (长度: {len(text)}字符, 使用术语库: {use_termbase}, 强制使用术语: {force_terms})")
        logger.info(f"原文预览: {text[:50]}..." if len(text) > 50 else f"原文: {text}")
        
        # 处理target_lang参数，修复bool类型错误
        if isinstance(target_lang, bool):
            logger.warning(f"目标语言参数为布尔值 ({target_lang})，使用默认值'en'")
            target_lang = "en"
        elif not isinstance(target_lang, str):
            logger.warning(f"目标语言参数类型错误 ({type(target_lang).__name__})，使用默认值'en'")
            target_lang = "en"
        elif target_lang.lower() in ["中文", "chinese", "zh", "zh-cn", "zh-tw"]:
            target_lang = "zh"
        elif target_lang.lower() in ["英文", "english", "en", "en-us", "en-gb"]:
            target_lang = "en"
            
        # 源语言参数也做类似处理
        if isinstance(source_lang, str):
            if source_lang.lower() in ["中文", "chinese", "zh", "zh-cn", "zh-tw"]:
                source_lang = "zh"
            elif source_lang.lower() in ["英文", "english", "en", "en-us", "en-gb"]:
                source_lang = "en"
        elif source_lang is not None and not isinstance(source_lang, str):
            logger.warning(f"源语言参数类型错误 ({type(source_lang).__name__})，使用自动检测")
            source_lang = None
        
        try:
            # 如果提供了匹配到的术语列表，并且要求强制使用术语
            if force_terms and matched_terms:
                logger.info(f"使用提供的术语列表进行翻译，共 {len(matched_terms)} 个术语")
                # 使用专用的强制术语翻译函数
                return self._translate_with_force_terms(text, matched_terms, source_lang, target_lang)
            
            # 调用AI引擎的翻译方法
            translated_text = self.ai_engine.translate_text(
                text=text,
                use_termbase=use_termbase,
                source_lang=source_lang,
                target_lang=target_lang
            )
            
            logger.info(f"翻译完成 (长度: {len(translated_text)}字符)")
            logger.info(f"译文预览: {translated_text[:50]}..." if len(translated_text) > 50 else f"译文: {translated_text}")
            
            return translated_text
            
        except Exception as e:
            logger.error(f"翻译过程发生错误: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return f"翻译出错: {str(e)}"
    
    def _retry_translation(self, text, source_lang, target_lang, matched_terms=None):
        """重新尝试翻译，强调目标语言和术语使用"""
        try:
            # 根据语言选择合适的提示词
            source_lang_name = "中文" if source_lang == "zh" else "英文" if source_lang == "en" else source_lang
            target_lang_name = "英文" if target_lang == "en" else "中文" if target_lang == "zh" else target_lang
            
            # 构建系统提示词
            system_prompt = f"""你是一位专业翻译。这是一个严格的翻译任务，你必须将文本从{source_lang_name}翻译成{target_lang_name}。
不允许返回原文，必须提供翻译结果。"""
            
            # 如果有匹配到的术语，添加到提示词中
            if matched_terms:
                system_prompt += "\n\n必须使用以下专业术语进行翻译:\n"
                for term in matched_terms:
                    system_prompt += f"- {term['source']} => {term['target']}\n"
                system_prompt += "\n术语必须精确匹配，不允许改变形式或用同义词替代。"
            
            system_prompt += "\n如果你不知道如何翻译某些专业术语，保持它们的原样并用括号标注。"
            
            user_prompt = f"""原文({source_lang_name}):{text}请提供准确的{target_lang_name}翻译:"""
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            translation = self.ai_engine.generate_response(messages)
            logger.info(f"重试翻译完成: {translation[:100]}")
            
            return self._clean_translation_result(translation)
            
        except Exception as e:
            logger.error(f"重试翻译时出错: {str(e)}")
            return None
    
    def _clean_translation_result(self, text):
        """清理翻译结果，删除不必要的角色前缀等"""
        if not text:
            return ""
        
        # 移除常见的角色前缀
        prefixes_to_remove = [
            "user:", "User:", "用户:", 
            "assistant:", "Assistant:", "助手:",
            "system:", "System:", "系统:",
            "Translation:", "翻译:", "翻译结果:"
        ]
        
        result = text
        for prefix in prefixes_to_remove:
            if result.startswith(prefix):
                result = result[len(prefix):].strip()
        
        # 删除结尾可能的引号和多余空白
        result = result.strip('" \t\n')
        
        return result
    
    def _get_term_context(self, text, term):
        """获取术语在文本中的上下文"""
        try:
            # 找到术语在文本中的位置
            pos = text.find(term)
            if pos == -1:
                return ""
            
            # 获取术语前后的文本（最多20个字符）
            start = max(0, pos - 20)
            end = min(len(text), pos + len(term) + 20)
            
            # 提取上下文
            context = text[start:end]
            if start > 0:
                context = "..." + context
            if end < len(text):
                context = context + "..."
            
            return context
        except:
            return ""
    
    def _translate_with_terms(self, text, matched_terms, source_lang, target_lang):
        """使用术语库进行翻译"""
        try:
            # 构建翻译提示
            source_lang_name = "中文" if source_lang == "zh" else "英文" if source_lang == "en" else source_lang
            target_lang_name = "英文" if target_lang == "en" else "中文" if target_lang == "zh" else target_lang
            
            # 加强系统提示，明确要求使用特定术语
            system_prompt = f"""你是一位专业翻译，需要将以下{source_lang_name}文本翻译成{target_lang_name}。
此任务的关键是：必须使用以下指定的专业术语进行翻译，不得改变这些术语的翻译方式。

专业术语对照表（必须严格按照下面的对应关系翻译）：
"""
            
            # 添加术语对照表，格式化突出显示
            for term in matched_terms:
                system_prompt += f"【{term['source']}】 => 【{term['target']}】\n"
                if term.get('context'):
                    system_prompt += f"  • 上下文: {term['context']}\n"
                if term.get('definition'):
                    system_prompt += f"  • 定义: {term['definition']}\n"
            
            system_prompt += """
翻译要求：
1. 必须使用上述专业术语表中的对应术语，这是最重要的要求
2. 保持专业风格和自然流畅
3. 保留原文的格式和标点
4. 只输出翻译结果，不要添加解释或前缀
5. 不要重复原文，只输出翻译

验证过程：
- 我将检查你的翻译是否正确使用了所有专业术语
- 如果未正确使用术语，翻译将被拒绝

"""
            
            # 用户提示更加简洁明确
            user_prompt = f"""需要翻译的{source_lang_name}文本：
{text}

请将上述文本翻译成{target_lang_name}，必须使用指定的专业术语："""
            
            # 构建消息列表
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            print(f"发送带有{len(matched_terms)}个专业术语的翻译请求...")
            logger.debug(f"专业术语翻译提示词: {system_prompt[:300]}...")
            
            # 执行翻译
            translation = self.ai_engine.generate_response(messages)
            logger.info("完成带术语的翻译")
            
            # 验证翻译结果中是否包含所有术语
            missing_terms = []
            for term in matched_terms:
                if term['target'] not in translation:
                    missing_terms.append(term)
                    logger.warning(f"警告：翻译结果中未找到术语 '{term['source']}' 的标准译法 '{term['target']}'")
            
            # 如果有缺失的术语，尝试再次翻译
            if missing_terms and len(missing_terms) < len(matched_terms):
                logger.info(f"翻译结果缺少 {len(missing_terms)} 个术语，尝试修正翻译...")
                
                # 更强力的修正提示，强调术语使用的重要性
                fix_prompt = f"""你的翻译需要修正。以下专业术语在翻译结果中不存在或使用不正确：

"""
                for term in missing_terms:
                    fix_prompt += f"【{term['source']}】必须翻译为【{term['target']}】\n"
                
                fix_prompt += f"""
原文：
{text}

你的翻译结果：
{translation}

请重新翻译，必须严格使用上述所有专业术语。这是一项精确的技术翻译任务，术语使用的准确性至关重要："""
                
                fix_messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": fix_prompt}
                ]
                
                # 执行修正翻译
                fixed_translation = self.ai_engine.generate_response(fix_messages)
                if fixed_translation != translation:
                    logger.info("已修正翻译结果")
                    translation = fixed_translation
            
            return translation
            
        except Exception as e:
            logger.error(f"带术语翻译时出错: {str(e)}")
            raise
    
    def _translate_without_terms(self, text, source_lang, target_lang):
        """不使用术语库的普通翻译"""
        try:
            # 根据语言选择合适的提示词
            source_lang_name = "中文" if source_lang == "zh" else "英文" if source_lang == "en" else source_lang
            target_lang_name = "英文" if target_lang == "en" else "中文" if target_lang == "zh" else target_lang
            
            system_prompt = f"""你是一位专业翻译，需要将{source_lang_name}文本翻译成{target_lang_name}。

翻译规则：
1. 保持专业风格和自然流畅
2. 保留原文的格式和标点
3. 只输出翻译结果，不要添加解释或前缀
4. 不要重复原文，只输出翻译
"""
            
            user_prompt = f"""需要翻译的{source_lang_name}文本:
{text}

请将上述文本翻译成{target_lang_name}:"""
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            translation = self.ai_engine.generate_response(messages)
            logger.info("完成普通翻译")
            
            # 确保结果不是原文
            if translation.strip() == text.strip():
                logger.warning("翻译结果与原文相同，重试翻译...")
                
                # 更强调的翻译提示
                retry_system_prompt = f"""你必须将以下{source_lang_name}文本翻译成{target_lang_name}。
这是一个严格的翻译任务，必须提供翻译结果，不得返回原文。"""
                
                retry_user_prompt = f"""原文({source_lang_name}):
{text}

翻译结果({target_lang_name}):"""
                
                retry_messages = [
                    {"role": "system", "content": retry_system_prompt},
                    {"role": "user", "content": retry_user_prompt}
                ]
                
                translation = self.ai_engine.generate_response(retry_messages)
            
            return translation
            
        except Exception as e:
            logger.error(f"普通翻译时出错: {str(e)}")
            raise
    
    def _analyze_format(self, text):
        """分析文本格式，返回格式信息"""
        format_info = {
            'paragraphs': [],  # 段落分隔位置
            'indentations': {},  # 每行的缩进
            'line_breaks': []  # 换行符位置
        }
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            # 记录缩进
            leading_spaces = len(line) - len(line.lstrip())
            format_info['indentations'][i] = leading_spaces
            
            # 记录空行（段落分隔）
            if line.strip() == "":
                format_info['paragraphs'].append(i)
            
        return format_info
    
    def _split_to_sentences(self, text):
        """将文本分割为句子，保留分隔符和格式"""
        # 使用更复杂的正则表达式来识别句子边界
        import re
        
        # 分段处理，保留段落结构
        paragraphs = text.split('\n')
        all_sentences = []
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                # 保留空行
                all_sentences.append("\n")
                continue
            
            # 中英文句子结束标记
            # 英文: .!?后跟空格或结束
            # 中文: 。！？后面可能没有空格
            pattern = r'(?<=[.!?。！？])\s*'
            
            # 分割句子但保留分隔符
            sentences = re.split(pattern, paragraph)
            
            # 处理结果，确保句尾有标点
            processed_sentences = []
            for i, sentence in enumerate(sentences):
                if not sentence.strip():
                    continue
                
                # 检查句子是否以标点结束
                if i < len(sentences) - 1:
                    # 查找最后一个句子终结符
                    end_chars = ['.', '!', '?', '。', '！', '？']
                    last_end_pos = max(sentence.rfind(c) for c in end_chars)
                    
                    if last_end_pos != -1:
                        # 在句子终结符后分割
                        processed_sentences.append(sentence[:last_end_pos+1])
                        if last_end_pos+1 < len(sentence):
                            remaining = sentence[last_end_pos+1:].strip()
                            if remaining:
                                processed_sentences.append(remaining)
                    else:
                        processed_sentences.append(sentence)
                else:
                    processed_sentences.append(sentence)
            
            # 将处理后的句子加入结果
            all_sentences.extend(processed_sentences)
            all_sentences.append("\n")  # 保留段落分隔
        
        return all_sentences
    
    def _rebuild_text(self, sentences, format_info):
        """根据原有格式重建文本"""
        # 重建段落结构
        result = []
        current_line = ""
        
        for sentence in sentences:
            if sentence == "\n":
                # 保留段落分隔
                result.append(current_line)
                result.append("")  # 空行
                current_line = ""
            elif sentence.strip():
                # 添加非空句子
                if current_line and not current_line.endswith(("\n", " ")):
                    current_line += " "  # 添加空格分隔符
                current_line += sentence
        
        # 添加最后一行
        if current_line:
            result.append(current_line)
        
        # 应用缩进
        for i, line in enumerate(result):
            if i in format_info['indentations']:
                indent = " " * format_info['indentations'][i]
                result[i] = indent + line
        
        # 连接所有行
        return "\n".join(result)
    
    def translate_file(self, file_path, use_termbase=False, target_lang="zh"):
        """翻译文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        # 获取文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 根据文件类型处理
        if file_ext == '.txt':
            return self.translate_text_file(file_path, use_termbase, target_lang)
        elif file_ext == '.md':
            return self.translate_markdown_file(file_path, use_termbase, target_lang)
        elif file_ext == '.docx':
            return self.translate_docx_file(file_path, use_termbase, target_lang)
        elif file_ext == '.pdf':
            return self.translate_pdf_file(file_path, use_termbase, target_lang)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")
    
    def translate_text_file(self, file_path, use_termbase, target_lang):
        """翻译文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # 翻译内容
        translated = self.translate(content, use_termbase, target_lang=target_lang)
        
        return translated
    
    def translate_markdown_file(self, file_path, use_termbase, target_lang):
        """翻译Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 将markdown分成代码块和非代码块
        blocks = self.split_markdown(content)
        
        translated_blocks = []
        for block in blocks:
            if block['type'] == 'code':
                # 代码块不翻译
                translated_blocks.append(block['content'])
            else:
                # 翻译非代码块
                translated = self.translate(
                    block['content'], 
                    use_termbase, 
                    target_lang=target_lang
                )
                translated_blocks.append(translated)
        
        # 合并翻译结果
        return ''.join(translated_blocks)
    
    def split_markdown(self, content):
        """将markdown分成代码块和非代码块"""
        # 查找代码块
        code_pattern = r'```[\s\S]*?```'
        code_blocks = re.findall(code_pattern, content)
        
        # 分割内容
        blocks = []
        remainder = content
        
        for code in code_blocks:
            parts = remainder.split(code, 1)
            
            if parts[0]:
                blocks.append({'type': 'text', 'content': parts[0]})
                
            blocks.append({'type': 'code', 'content': code})
            
            remainder = parts[1] if len(parts) > 1 else ""
        
        if remainder:
            blocks.append({'type': 'text', 'content': remainder})
            
        return blocks
    
    def translate_docx_file(self, file_path, use_termbase, target_lang):
        """翻译Word文档"""
        doc = Document(file_path)
        
        # 创建新文档
        new_doc = Document()
        
        # 翻译段落
        for para in doc.paragraphs:
            if para.text.strip():
                translated = self.translate(
                    para.text, 
                    use_termbase, 
                    target_lang=target_lang
                )
                new_doc.add_paragraph(translated)
            else:
                new_doc.add_paragraph()
        
        # 翻译表格
        for table in doc.tables:
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text.strip():
                        translated = self.translate(
                            cell.text, 
                            use_termbase, 
                            target_lang=target_lang
                        )
                        new_table.cell(i, j).text = translated
        
        # 获取输出文件路径
        output_path = os.path.splitext(file_path)[0] + f"_{target_lang}.docx"
        
        # 保存翻译后的文档
        new_doc.save(output_path)
        
        return output_path
    
    def translate_pdf_file(self, file_path, use_termbase, target_lang):
        """翻译PDF文件"""
        # 读取PDF内容
        reader = PdfReader(file_path)
        content = ""
        
        for page in reader.pages:
            content += page.extract_text() + "\n\n"
        
        # 翻译内容
        translated = self.translate(content, use_termbase, target_lang=target_lang)
        
        # 获取输出文件路径
        output_path = os.path.splitext(file_path)[0] + f"_{target_lang}.txt"
        
        # 保存翻译后的内容
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated)
        
        return output_path
    
    def translate_selected_text(self, text, target_lang="zh"):
        """翻译选中的文本"""
        if not text or len(text) > 5000:  # 限制长度
            return ""
            
        # 使用术语库，但简化处理
        use_termbase = self.settings.get('use_term_base', False)
        
        return self.translate(text, use_termbase, target_lang=target_lang)
    
    def process_terminology(self, source_text, source_lang="zh", target_lang="en"):
        """处理术语库，识别并替换术语"""
        if not hasattr(self, 'term_loader') or not self.term_loader:
            print("[WARNING] 术语库未初始化，跳过术语替换")
            return source_text
        
        try:
            # 加载术语
            terms = self.term_loader.terms
            if not terms:
                return source_text
            
            # 按长度排序术语，优先替换较长的术语
            sorted_terms = sorted(terms.keys(), key=len, reverse=True)
            
            # 替换文本
            result_text = source_text
            replacements = []
            
            for term in sorted_terms:
                if term in result_text:
                    term_data = terms[term]
                    target_term = term_data.get('target_term', term_data.get('definition', ''))
                    
                    # 检查语言是否匹配
                    metadata = term_data.get('metadata', {})
                    src_lang = metadata.get('source_lang', 'zh')
                    tgt_lang = metadata.get('target_lang', 'en')
                    
                    if src_lang == source_lang and tgt_lang == target_lang and target_term:
                        # 记录替换
                        replacements.append({
                            'source': term,
                            'target': target_term
                        })
            
            # 记录术语替换信息
            if replacements:
                self.last_term_replacements = replacements
                print(f"[INFO] 发现 {len(replacements)} 个术语匹配")
            
            return result_text
        
        except Exception as e:
            print(f"[ERROR] 术语替换失败: {e}")
            import traceback
            traceback.print_exc()
            return source_text 
    
    def get_translation_prompt(self, text, source_lang="auto", target_lang="zh"):
        """生成翻译提示词"""
        source_lang_map = {
            "auto": "自动检测",
            "zh": "中文",
            "en": "英文",
            "ja": "日文",
            "ko": "韩文",
            "fr": "法文",
            "de": "德文",
            "es": "西班牙文",
            "ru": "俄文"
        }
        
        target_lang_map = {
            "zh": "中文",
            "en": "英文",
            "ja": "日文",
            "ko": "韩文",
            "fr": "法文",
            "de": "德文",
            "es": "西班牙文",
            "ru": "俄文"
        }
        
        source_lang_name = source_lang_map.get(source_lang, "自动检测")
        target_lang_name = target_lang_map.get(target_lang, "中文")
        
        # 优化提示词，明确源语言和目标语言
        if target_lang == "en":
            prompt = f"""You are a professional translator. Please translate the following text from {source_lang_name} to English. 
Maintain the original meaning, professional terminology, and formatting.

Text to translate:
{text}

English Translation:
"""
        else:
            prompt = f"""你是一位专业翻译，请将以下{source_lang_name}文本翻译成{target_lang_name}。
保持原意、专业性和格式，进行逐行翻译。

原文:
{text}

翻译结果:
"""
        
        return prompt 

    def _fix_missing_terms(self, translation, original_text, matched_terms, source_lang, target_lang):
        """修正翻译结果中缺失的术语"""
        try:
            # 构建更强调术语使用的提示
            source_lang_name = "中文" if source_lang == "zh" else "英文" if source_lang == "en" else source_lang
            target_lang_name = "英文" if target_lang == "en" else "中文" if target_lang == "zh" else target_lang
            
            system_prompt = f"""你是一位专业翻译，需要修正已翻译的文本，确保正确使用专业术语。
            
必须使用的专业术语对照表：
"""
            
            # 添加术语对照表
            for term in matched_terms:
                system_prompt += f"- {term['source']} => {term['target']}\n"
            
            user_prompt = f"""原文({source_lang_name}):
{original_text}

当前的翻译({target_lang_name}):
{translation}

请修正上述翻译，必须严格使用给定的专业术语。
修正后的翻译(只需要给出修正后的翻译结果):"""
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            fixed_translation = self.ai_engine.generate_response(messages)
            logger.debug(f"修正后的翻译: {fixed_translation}")
            
            return self._clean_translation_result(fixed_translation)
        except Exception as e:
            logger.error(f"修正术语时出错: {str(e)}")
            return translation 

    def _translate_with_force_terms(self, text: str, matched_terms: List, source_lang: Optional[str], target_lang: str) -> str:
        """使用强制术语进行翻译，先替换术语再翻译"""
        try:
            # 1. 先进行术语替换
            replaced_text = text
            for term in matched_terms:
                replaced_text = replaced_text.replace(term['source'], term['target'])
            
            # 2. 构建更严格的翻译提示
            prompt = f"""请将以下文本严格翻译成{target_lang}，必须使用以下术语：
术语表：
"""
            # 添加术语表
            for term in matched_terms:
                prompt += f"- {term['source']} => {term['target']}\n"
            
            prompt += f"""
翻译要求：
1. 严格使用上述术语表中的术语
2. 只输出翻译结果，不要添加任何解释或额外信息
3. 保持原文结构和格式

要翻译的文本：
{replaced_text}

翻译结果："""
            
            messages = [
                {"role": "system", "content": prompt},
                {"role": "user", "content": replaced_text}
            ]
            
            # 3. 执行翻译
            translation = self.ai_engine.generate_response(messages)
            
            # 4. 清理结果并返回
            return self._clean_translation_result(translation)
            
        except Exception as e:
            logger.error(f"强制术语翻译出错: {e}")
            return f"翻译出错: {str(e)}"

    def _find_matching_terms_in_text(self, text: str, terms: Dict, source_lang: str = None, target_lang: str = None):
        """在文本中寻找匹配的术语，返回匹配到的术语列表"""
        logger.info(f"开始术语匹配，文本长度: {len(text)}，术语库大小: {len(terms)}")
        
        # 检查输入
        if not text or not terms:
            logger.warning("输入文本或术语库为空，无法进行匹配")
            return []
        
        # 如果未指定源语言，尝试检测
        if not source_lang:
            try:
                source_lang = detect(text)
                logger.info(f"自动检测到源语言: {source_lang}")
            except:
                source_lang = "zh"  # 默认中文
                logger.info(f"语言检测失败，默认使用源语言: {source_lang}")
        
        # 标准化语言代码
        if source_lang and source_lang.lower() in ['zh-cn', 'zh-tw', 'chinese', '中文']:
            source_lang = 'zh'
        elif source_lang and source_lang.lower() in ['en-us', 'en-gb', 'english', '英文']:
            source_lang = 'en'
        
        if target_lang and target_lang.lower() in ['zh-cn', 'zh-tw', 'chinese', '中文']:
            target_lang = 'zh'
        elif target_lang and target_lang.lower() in ['en-us', 'en-gb', 'english', '英文']:
            target_lang = 'en'
        
        matched_terms = []
        
        # 排序术语列表，优先匹配长术语（避免子字符串问题）
        sorted_term_items = sorted(
            [(k, v) for k, v in terms.items()],
            key=lambda x: len(x[0]),
            reverse=True
        )
        
        for term_key, term_info in sorted_term_items:
            # 获取源术语和目标术语
            source_term = term_info.get('source_term', term_key)
            target_term = term_info.get('target_term', '')
            
            # 跳过无效术语
            if not source_term or not target_term:
                continue
            
            # 简化匹配逻辑：如果源术语在文本中，则认为匹配成功
            if source_term in text:
                logger.info(f"匹配到术语: '{source_term}' => '{target_term}'")
                matched_terms.append({
                    'source': source_term,
                    'target': target_term,
                    'definition': term_info.get('definition', '')
                })
        
        logger.info(f"术语匹配完成，共匹配到 {len(matched_terms)} 个术语")
        return matched_terms 

    def _build_translation_prompt(self, text, use_termbase, matched_terms, target_lang):
        """构建翻译提示词"""
        # 确定源语言和目标语言名称
        source_lang = None
        try:
            detected_lang = detect(text)
            source_lang = 'zh' if detected_lang.lower() in ['zh-cn', 'zh-tw', 'zh'] else 'en' if detected_lang.lower() == 'en' else detected_lang
        except:
            source_lang = 'zh' if any(ord(c) > 127 for c in text) else 'en'
        
        source_lang_name = "中文" if source_lang == "zh" else "英文" if source_lang == "en" else source_lang
        target_lang_name = "中文" if target_lang == "zh" else "英文" if target_lang == "en" else target_lang
        
        logger.info(f"翻译方向: {source_lang_name} -> {target_lang_name}")
        
        # 基础提示词
        prompt = (
            f"请将以下{source_lang_name}文本翻译成{target_lang_name}。\n\n"
            "翻译要求：\n"
            "1. 保持原文的结构和格式\n"
            "2. 确保翻译准确、专业且流畅\n"
        )
        
        # 如果有匹配到术语，添加术语表
        if use_termbase and matched_terms:
            prompt += "3. 准确使用以下专业术语表：\n\n专业术语表：\n"
            
            for i, term in enumerate(matched_terms, 1):
                prompt += f"{i}. {term['source']} → {term['target']}"
                if term.get('definition'):
                    prompt += f" ({term['definition']})"
                prompt += "\n"
            
            # 添加术语使用指导，但保持简洁
            prompt += "\n翻译中应准确使用以上术语，保持专业性的同时确保整体流畅自然。\n"
        
        # 添加原文
        prompt += f"\n原文 ({source_lang_name})：\n{text}\n\n"
        prompt += f"译文 ({target_lang_name})：\n"
        
        logger.info(f"翻译提示构建完成，长度: {len(prompt)}")
        return prompt 

    def _verify_and_fix_term_usage(self, translation, matched_terms, original_text):
        """验证并修正翻译中的术语使用"""
        if not matched_terms:
            return translation
        
        logger.info(f"开始验证术语使用，需验证 {len(matched_terms)} 个术语")
        
        # 检查每个术语是否在翻译结果中出现
        missing_terms = []
        for term in matched_terms:
            if term['source'] in original_text and term['target'] not in translation:
                missing_terms.append(term)
                logger.warning(f"术语缺失: '{term['source']}' 应译为 '{term['target']}'")
        
        # 如果没有缺失术语，直接返回原翻译
        if not missing_terms:
            logger.info("所有术语均已正确使用")
            return translation
        
        # 如果缺失术语较少，尝试直接修复翻译
        if len(missing_terms) <= 3:
            logger.info(f"尝试修复 {len(missing_terms)} 个缺失术语")
            
            # 构建修复提示
            fix_prompt = f"""请修正以下翻译，确保正确使用专业术语:

原文:
{original_text}

当前翻译:
{translation}

需要修正的术语:
"""
            for term in missing_terms:
                fix_prompt += f"- '{term['source']}' 必须翻译为 '{term['target']}'\n"
            
            fix_prompt += "\n请提供修正后的完整翻译:"
            
            # 创建消息列表
            messages = [
                {"role": "system", "content": "你是一位专业翻译修正专家。请修正翻译中的术语使用问题，保持其他内容不变。"},
                {"role": "user", "content": fix_prompt}
            ]
            
            # 生成修正后的翻译
            try:
                fixed_translation = self.ai_engine.generate_response(messages)
                logger.info(f"修正后的翻译长度: {len(fixed_translation)}")
                return fixed_translation
            except Exception as e:
                logger.error(f"修正翻译出错: {e}")
        
        # 如果缺失术语较多或修复失败，添加警告信息
        logger.warning(f"存在 {len(missing_terms)} 个术语使用问题无法自动修复")
        warning = "\n\n[术语翻译参考]\n"
        for term in missing_terms:
            warning += f"- {term['source']} → {term['target']}\n"
        
        return translation + warning 